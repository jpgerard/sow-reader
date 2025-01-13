import re
import pandas as pd
import spacy
import pdfplumber
from pathlib import Path
from typing import List, Dict
from concurrent.futures import ThreadPoolExecutor
from docx import Document

# Load spaCy model
try:
    nlp = spacy.load("en_core_web_sm")
except OSError:
    import subprocess
    subprocess.check_call(["python", "-m", "spacy", "download", "en_core_web_sm"])
    nlp = spacy.load("en_core_web_sm")

class SOWProcessor:
    """Processes Statement of Work documents to extract requirements."""
    def __init__(self):
        """Initialize the SOW processor."""
        self.categories = {
            'Technical': r'\b(?:technical|system|software|hardware|network|infrastructure|security)\b',
            'Process': r'\b(?:process|procedure|workflow|method|approach|implementation)\b',
            'Service': r'\b(?:service|support|maintenance|operation|performance)\b',
            'Documentation': r'\b(?:document|report|plan|deliverable|submission)\b',
            'Compliance': r'\b(?:comply|compliance|requirement|standard|regulation|policy)\b',
            'Training': r'\b(?:training|instruction|education|knowledge|skill)\b'
        }
        self.mandatory_keywords = ["shall", "must", "required to", "responsible for", "directed to", "required"]
        self.informative_keywords = ["will", "plans to", "intends to", "should", "expected to", "recommended"]
        self.action_verbs = ["provide", "implement", "support", "develop", "maintain"]

    def process_document(self, file_path: str) -> List[Dict]:
        """Process a document and extract requirements."""
        text = self._load_document(file_path)
        return self.extract_requirements_from_text(text)

    def _load_document(self, file_path: str) -> str:
        """Load and clean text from a document."""
        file_path = Path(file_path)
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        suffix = file_path.suffix.lower()
        if suffix == '.pdf':
            text_chunks = []
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    try:
                        # Extract words directly and reconstruct text
                        words = page.extract_words()
                        if words:
                            # Sort words by top position and x position
                            words.sort(key=lambda w: (w['top'], w['x0']))
                            current_line = []
                            current_top = words[0]['top']
                            
                            for word in words:
                                if abs(word['top'] - current_top) > 3:  # New line
                                    if current_line:
                                        line_text = ' '.join(current_line)
                                        if not re.search(r'(?:Source|Page|For Official Use Only|^\d+$)', line_text):
                                            text_chunks.append(line_text)
                                    current_line = []
                                    current_top = word['top']
                                current_line.append(word['text'])
                            
                            # Add the last line
                            if current_line:
                                line_text = ' '.join(current_line)
                                if not re.search(r'(?:Source|Page|For Official Use Only|^\d+$)', line_text):
                                    text_chunks.append(line_text)
                    except Exception as e:
                        print(f"Error processing page: {e}")
                        continue
            return '\n'.join(text_chunks)
        elif suffix == '.docx':
            doc = Document(file_path)
            paragraphs = []
            
            # Extract text from paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = ' | '.join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        paragraphs.append(row_text)
            
            return '\n'.join(paragraphs)
        else:
            raise ValueError("Unsupported file format. Only PDF and DOCX files are supported.")

    def extract_requirements_from_text(self, text: str) -> List[Dict]:
        """Extract requirements from text content."""
        sections = self._parse_sections(text)
        all_requirements = []
        for section in sections:
            requirements = self._extract_requirements(section)
            all_requirements.extend(requirements)
        unique_requirements = self._deduplicate_requirements(all_requirements)
        categorized_requirements = self._categorize_requirements(unique_requirements)
        return self._sort_requirements(categorized_requirements)

    def _parse_sections(self, text: str) -> List[Dict]:
        """Extract sections using regex with flexible patterns."""
        sections = []
        # Split text into lines for better section matching
        lines = text.split('\n')
        current_section = None
        current_content = []
        
        for line in lines:
            # Try to match section headers
            section_match = re.match(r'^([A-Z](?:\.\d+)?|\d+\.\d+(?:\.\d+)?)\s+(.+)$', line.strip())
            if section_match:
                # Save previous section if exists
                if current_section and current_content:
                    sections.append({
                        'id': current_section,
                        'content': ' '.join(current_content)
                    })
                # Start new section
                current_section = section_match.group(1)
                current_content = [section_match.group(2)]
            elif current_section and line.strip():
                current_content.append(line.strip())
        
        # Add the last section
        if current_section and current_content:
            sections.append({
                'id': current_section,
                'content': ' '.join(current_content)
            })
        
        return sections

    def _extract_requirements(self, section: Dict) -> List[Dict]:
        """Extract requirements from a section."""
        sentences = re.split(r'(?<=[.!?])\s+(?=[A-Z])', section['content'])
        requirements = []
        for sentence in sentences:
            sentence = sentence.strip()
            if len(sentence.split()) < 5:
                continue
            result = self._analyze_requirement(sentence)
            if result['is_requirement']:
                requirements.append({
                    'section_id': section['id'],
                    'text': sentence,
                    'type': result['type'],
                    'confidence': result['confidence']
                })
        return requirements

    def _analyze_requirement(self, sentence: str) -> Dict:
        """Analyze if a sentence is a requirement."""
        is_mandatory = any(re.search(rf'\b{kw}\b', sentence, re.IGNORECASE) for kw in self.mandatory_keywords)
        is_informative = any(re.search(rf'\b{kw}\b', sentence, re.IGNORECASE) for kw in self.informative_keywords)
        has_action = any(verb in sentence.lower() for verb in self.action_verbs)
        return {
            'is_requirement': is_mandatory or is_informative,
            'type': 'Mandatory' if is_mandatory else ('Informative' if is_informative else None),
            'confidence': 0.8 if is_mandatory else (0.6 if is_informative else 0.0) + (0.1 if has_action else 0.0)
        }

    def _deduplicate_requirements(self, requirements: List[Dict]) -> List[Dict]:
        """Remove duplicate requirements."""
        unique_requirements = {}
        for req in requirements:
            normalized_text = re.sub(r'\s+', ' ', req['text'].lower().strip())
            normalized_text = re.sub(r'[^\w\s]', '', normalized_text)
            if normalized_text not in unique_requirements or req['confidence'] > unique_requirements[normalized_text]['confidence']:
                unique_requirements[normalized_text] = req
        return list(unique_requirements.values())

    def _categorize_requirements(self, requirements: List[Dict]) -> List[Dict]:
        """Categorize requirements."""
        for req in requirements:
            for category, pattern in self.categories.items():
                if re.search(pattern, req['text'], re.IGNORECASE):
                    req['category'] = category
                    break
            else:
                req['category'] = 'General'
        return requirements

    def _sort_requirements(self, requirements: List[Dict]) -> List[Dict]:
        """Sort requirements by section ID."""
        def sort_key(req):
            # Split section ID into components for natural sorting
            section_id = req['section_id']
            parts = re.split(r'[\.\s]', section_id)
            key_parts = []
            for part in parts:
                try:
                    key_parts.append(int(part))
                except ValueError:
                    key_parts.append(part)
            # Sort by section ID first, then by confidence (descending), then by text
            return (key_parts, -req['confidence'], req['text'])
            
        return sorted(requirements, key=sort_key)

    def save_requirements_to_csv(self, requirements: List[Dict], filename: str = "extracted_requirements.csv"):
        """Save requirements to a CSV file."""
        df = pd.DataFrame(requirements)
        df.to_csv(filename, index=False)
